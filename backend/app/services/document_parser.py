from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.models.document import DocumentParseStatus

MIN_EXTRACTED_CHARACTERS = 30


@dataclass(frozen=True)
class ParsedDocument:
    extracted_text: str | None
    parse_status: DocumentParseStatus
    metadata: dict = field(default_factory=dict)


def extract_document_text(data: bytes, filename: str, content_type: str | None = None) -> ParsedDocument:
    extension = Path(filename).suffix.lower().lstrip(".")
    metadata = {
        "file_type": extension or content_type or "unknown",
        "character_count": 0,
    }

    try:
        if extension == "txt":
            text = _extract_txt(data)
        elif extension == "pdf":
            text, page_count = _extract_pdf(data)
            metadata["page_count"] = page_count
        elif extension == "docx":
            text = _extract_docx(data)
        else:
            return ParsedDocument(
                extracted_text=None,
                parse_status=DocumentParseStatus.FAILED,
                metadata={**metadata, "error": "unsupported_file_type"},
            )
    except Exception as exc:
        return ParsedDocument(
            extracted_text=None,
            parse_status=DocumentParseStatus.FAILED,
            metadata={**metadata, "error": type(exc).__name__},
        )

    normalized = _normalize_text(text)
    metadata["character_count"] = len(normalized)
    if len(normalized) < MIN_EXTRACTED_CHARACTERS:
        return ParsedDocument(
            extracted_text=normalized or None,
            parse_status=DocumentParseStatus.OCR_REQUIRED,
            metadata=metadata,
        )

    return ParsedDocument(
        extracted_text=normalized,
        parse_status=DocumentParseStatus.PARSED,
        metadata=metadata,
    )


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_pdf(data: bytes) -> tuple[str, int]:
    reader = PdfReader(BytesIO(data))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages), len(reader.pages)


def _extract_docx(data: bytes) -> str:
    try:
        document = DocxDocument(BytesIO(data))
    except BadZipFile:
        raise
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _normalize_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()
